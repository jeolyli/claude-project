"""
Convert docx to PDF using python-docx + reportlab.
Preserves: headings (H1/H2), body text, tables, and basic formatting.
"""
import sys, os
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

INPUT = sys.argv[1] if len(sys.argv) > 1 else "灵记_项目运营手册.docx"
OUTPUT = INPUT.replace(".docx", ".pdf")

# ===== Register Chinese font =====
# Try to find a Chinese font available on macOS
FONT_PATHS = [
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    "/Library/Fonts/Arial Unicode.ttf",
]

FONT_NAME = "PingFang"
font_path = None
for fp in FONT_PATHS:
    if os.path.exists(fp):
        font_path = fp
        break

if font_path:
    pdfmetrics.registerFont(TTFont(FONT_NAME, font_path))
else:
    print("Warning: No Chinese font found. Chinese characters may not render.")
    FONT_NAME = "Helvetica"

# ===== Styles =====
styles = getSampleStyleSheet()
PRIMARY = HexColor("#FF85A2")
DARK = HexColor("#3D2C33")
GREY = HexColor("#757575")
LIGHT = HexColor("#FFF5F7")
BORDER = HexColor("#FFD4DE")

body_style = ParagraphStyle(
    "BodyCN", parent=styles["Normal"],
    fontName=FONT_NAME, fontSize=10, leading=16, textColor=DARK,
    spaceBefore=4, spaceAfter=4,
)

h1_style = ParagraphStyle(
    "H1CN", parent=styles["Heading1"],
    fontName=FONT_NAME, fontSize=18, leading=24, textColor=PRIMARY,
    spaceBefore=20, spaceAfter=10, bold=True,
)

h2_style = ParagraphStyle(
    "H2CN", parent=styles["Heading2"],
    fontName=FONT_NAME, fontSize=14, leading=20, textColor=DARK,
    spaceBefore=14, spaceAfter=8, bold=True,
)

note_style = ParagraphStyle(
    "NoteCN", parent=styles["Normal"],
    fontName=FONT_NAME, fontSize=9, leading=14, textColor=GREY,
    leftIndent=12, spaceBefore=2, spaceAfter=2,
)

center_style = ParagraphStyle(
    "CenterCN", parent=styles["Normal"],
    fontName=FONT_NAME, fontSize=10, leading=16, textColor=DARK,
    alignment=TA_CENTER, spaceBefore=2, spaceAfter=2,
)

cover_title = ParagraphStyle(
    "CoverTitle", parent=styles["Normal"],
    fontName=FONT_NAME, fontSize=28, leading=36, textColor=PRIMARY,
    alignment=TA_CENTER, spaceBefore=6, spaceAfter=4, bold=True,
)

cover_subtitle = ParagraphStyle(
    "CoverSub", parent=styles["Normal"],
    fontName=FONT_NAME, fontSize=18, leading=24, textColor=DARK,
    alignment=TA_CENTER, spaceBefore=4, spaceAfter=4,
)

# ===== Read DOCX =====
doc = Document(INPUT)
story = []

# Track if we've already seen content to avoid duplicate cover
title_seen = False

for element in doc.element.body:
    tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

    if tag == "p":
        # Get paragraph
        para = None
        for p in doc.paragraphs:
            if p._element is element:
                para = p
                break
        if para is None:
            continue

        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue

        style_name = para.style.name if para.style else "Normal"

        # Determine heading level
        if style_name.startswith("Heading 1") or style_name == "Heading1":
            story.append(Paragraph(text, h1_style))
        elif style_name.startswith("Heading 2") or style_name == "Heading2":
            story.append(Paragraph(text, h2_style))
        elif style_name == "NoteCN" or para.paragraph_format.left_indent:
            story.append(Paragraph(text, note_style))
        elif para.alignment == TA_CENTER:
            story.append(Paragraph(text, center_style))
        else:
            story.append(Paragraph(text, body_style))

    elif tag == "tbl":
        # Process table
        for table in doc.tables:
            if table._element is element:
                rows_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(Paragraph(cell.text.strip(), body_style))
                    rows_data.append(row_data)

                if rows_data:
                    col_count = len(rows_data[0])
                    col_widths = [460 / col_count] * col_count

                    tbl = Table(rows_data, colWidths=col_widths)
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), LIGHT),
                        ("TEXTCOLOR", (0, 0), (-1, 0), PRIMARY),
                        ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 8))
                break

# ===== Build PDF =====
pdf_doc = SimpleDocTemplate(
    OUTPUT,
    pagesize=A4,
    topMargin=25 * mm,
    bottomMargin=25 * mm,
    leftMargin=25 * mm,
    rightMargin=25 * mm,
    title="灵记 (LingJi) — 项目运营手册",
    author="灵记",
)
pdf_doc.build(story)
print(f"✅ PDF generated: {OUTPUT}")
